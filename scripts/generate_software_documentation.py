from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_FILE = Path(__file__).resolve().parents[1] / "Software_Documentation.docx"
TODAY = date(2026, 4, 14)
COMPANY_NAME = "CertifyPro"
SYSTEM_NAME = "CA Certificate Management System"


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)

    document.core_properties.title = "Software System Documentation"
    document.core_properties.subject = "Certificate Management System"
    document.core_properties.author = COMPANY_NAME
    document.core_properties.company = COMPANY_NAME


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_header_footer(section) -> None:
    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.text = f"{COMPANY_NAME} | {SYSTEM_NAME}"
    if header_para.runs:
        header_para.runs[0].font.size = Pt(9)
        header_para.runs[0].font.bold = True
        header_para.runs[0].font.color.rgb = RGBColor(68, 68, 68)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("Page ")
    add_page_number(footer_para)
    for run in footer_para.runs:
        run.font.size = Pt(9)


def add_cover_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run(COMPANY_NAME)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    p.space_before = Pt(140)
    run = p.add_run("Software System Documentation")
    run.font.size = Pt(28)
    run.font.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Certificate Management System")
    run.font.size = Pt(16)
    run.italic = True
    run.font.color.rgb = RGBColor(68, 68, 68)

    document.add_paragraph()

    details = [
        ("Version", "1.0"),
        ("Date", TODAY.strftime("%d %B %Y")),
        ("Prepared For", "Software Inspection Team"),
    ]

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in details:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(40)
    run = p.add_run(
        "Prepared as a formal overview of the certificate application, controls, "
        "workflow, access model, and inspection-ready operational considerations."
    )
    run.font.size = Pt(11)

    document.add_page_break()


def add_toc(document: Document) -> None:
    document.add_heading("Table of Contents", level=1)
    p = document.add_paragraph()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Update the table of contents in Word if page numbers do not refresh automatically."

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    p._p.append(fld_char_begin)
    p._p.append(instr_text)
    p._p.append(fld_char_separate)
    p._p.append(placeholder)
    p._p.append(fld_char_end)
    document.add_page_break()


def add_bullets(document: Document, items) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_system_overview(document: Document) -> None:
    document.add_heading("System Overview", level=1)
    document.add_paragraph(
        f"{SYSTEM_NAME} is a web-based business application used to create, manage, "
        "store, review, update, and export professional certificates in a controlled manner. "
        "The system supports certificate operations for authorized users, centralizes record "
        "keeping, and maintains audit visibility for inspection and administrative review."
    )

    document.add_heading("Key Features and Functionalities", level=2)
    add_bullets(
        document,
        [
            "User authentication with role-aware access handling for administrators and operational users.",
            "Certificate creation, listing, retrieval, update, and deletion through standardized API workflows.",
            "Support for multiple certificate categories such as turnover, net worth, utilisation, RERA, NBFC, and list-of-directors certificates.",
            "Persistent certificate storage with structured payload handling and record timestamps.",
            "History and audit tracking of key actions such as login, certificate creation, update, and deletion.",
            "Temporary access issuance and revocation for controlled operational use.",
            "DOCX export support for supported certificate outputs.",
            "Operational access controls for office-based or geolocation-based validation where applicable.",
        ],
    )

    document.add_heading("Technology Stack Used", level=2)
    stack_table = document.add_table(rows=1, cols=2)
    stack_table.style = "Table Grid"
    stack_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stack_table.rows[0].cells[0].text = "Layer"
    stack_table.rows[0].cells[1].text = "Technology"
    for left, right in [
        ("Frontend", "React with Vite, React Router, Axios"),
        ("Backend", "Python, FastAPI, Uvicorn"),
        ("Database Layer", "SQLAlchemy ORM"),
        ("Database", "PostgreSQL / Neon-compatible deployment"),
        ("Document Generation", "python-docx"),
        ("Deployment", "Render-based web service hosting"),
    ]:
        row = stack_table.add_row().cells
        row[0].text = left
        row[1].text = right


def add_how_it_works(document: Document) -> None:
    document.add_heading("How the Software Works", level=1)

    document.add_heading("Step-by-Step Workflow", level=2)
    steps = [
        "The authorized user opens the frontend application and signs in with valid credentials.",
        "The frontend submits the login request to the backend authentication endpoint.",
        "The backend validates the credentials, creates a session token, and returns the authenticated user profile.",
        "The user opens the required certificate form and provides business, entity, and certificate-specific details.",
        "The backend validates the submitted payload and checks role-based permissions before processing the request.",
        "Validated certificate data is stored in the database with certificate metadata, timestamps, and ownership details.",
        "The system records key activities in the history log for operational auditability.",
        "Users can later review certificates, update approved records, delete permitted records, or export supported documents.",
    ]
    for idx, step in enumerate(steps, start=1):
        document.add_paragraph(f"{idx}. {step}")

    document.add_heading("Input → Process → Output Flow", level=2)
    flow_table = document.add_table(rows=1, cols=3)
    flow_table.style = "Table Grid"
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = flow_table.rows[0].cells
    headers[0].text = "Input"
    headers[1].text = "Process"
    headers[2].text = "Output"
    flows = [
        (
            "Login credentials and access request",
            "Credential verification, token/session creation, access check",
            "Authenticated user session",
        ),
        (
            "Certificate form details",
            "Validation, category rules, record persistence",
            "Stored certificate record with unique ID",
        ),
        (
            "Update or edit request",
            "Ownership/role validation, payload merge, save",
            "Updated certificate record",
        ),
        (
            "Deletion request",
            "Authorization and ownership validation, removal, audit log",
            "Deleted record confirmation",
        ),
        (
            "Export request",
            "Data fetch, document composition, file streaming",
            "Downloadable document file",
        ),
    ]
    for input_text, process_text, output_text in flows:
        row = flow_table.add_row().cells
        row[0].text = input_text
        row[1].text = process_text
        row[2].text = output_text

    document.add_heading("System Architecture Overview", level=2)
    add_bullets(
        document,
        [
            "Presentation Layer: React-based web interface used by administrators and operational users.",
            "API Layer: FastAPI service responsible for authentication, business validation, certificate APIs, and audit workflows.",
            "Persistence Layer: SQLAlchemy-managed relational data layer storing users, sessions, certificates, history, and office metadata.",
            "Document Layer: DOCX generation utilities for supported exportable certificate outputs.",
            "Hosting Layer: Cloud-hosted deployment with external database connectivity and environment-based configuration.",
        ],
    )


def add_roles(document: Document) -> None:
    document.add_heading("Person Roles & Responsibilities", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Role"
    hdr[1].text = "Responsibilities"
    hdr[2].text = "Location/Department"

    rows = [
        (
            "Project Manager",
            "Coordinates scope, delivery planning, stakeholder communication, and inspection readiness.",
            "Project Management Office / Head Office",
        ),
        (
            "Software Developer",
            "Designs, develops, integrates, and maintains frontend, backend, and database components.",
            "Engineering / Application Development",
        ),
        (
            "QA Tester",
            "Validates workflows, verifies certificate outputs, regression tests releases, and records defects.",
            "Quality Assurance / Testing",
        ),
        (
            "System Admin",
            "Manages deployment settings, environment configuration, backups, service monitoring, and production access.",
            "IT Operations / Infrastructure",
        ),
        (
            "End User",
            "Uses the application to create, review, and manage certificate records according to assigned permissions.",
            "Operations / Certificate Processing",
        ),
    ]
    for role, responsibility, location in rows:
        row = table.add_row().cells
        row[0].text = role
        row[1].text = responsibility
        row[2].text = location


def add_access_levels(document: Document) -> None:
    document.add_heading("User Access Levels", level=1)
    document.add_paragraph(
        "The live application enforces administrator and operational user controls, while the following "
        "matrix documents the broader governance model used for inspection and business review."
    )

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = "Role"
    headers[1].text = "Authentication Scope"
    headers[2].text = "Certificate Access"
    headers[3].text = "Administrative Access"
    headers[4].text = "Typical Usage"

    data = [
        (
            "Admin",
            "Full login access with persistent credentials",
            "Create, view, update, delete all permitted records",
            "Manage users, temporary access, office settings, and oversight activities",
            "System owner or authorized administrator",
        ),
        (
            "Manager",
            "Authorized supervisory access",
            "Review team outputs, inspect records, approve operational completeness",
            "Limited configuration visibility as approved by governance",
            "Department lead or compliance reviewer",
        ),
        (
            "Staff",
            "Operational access with validated credentials",
            "Create and manage assigned certificate records within permission limits",
            "No platform-wide administration",
            "Certificate operations or processing staff",
        ),
        (
            "Read-Only",
            "Controlled access for observation or audit",
            "View approved data and reports without modification rights",
            "No configuration or transaction rights",
            "Inspection, audit, or reporting stakeholders",
        ),
    ]
    for row_data in data:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value


def add_inspection_checklist(document: Document) -> None:
    document.add_heading("Software Inspection Checklist", level=1)
    checklist_items = [
        "Verify that the application starts successfully and the health endpoint responds as expected.",
        "Confirm that user authentication is functioning and invalid credentials are rejected.",
        "Verify that administrator access can manage system operations and that non-admin access is role-restricted.",
        "Check that certificate data entry forms validate mandatory fields and category-specific rules.",
        "Confirm that certificate records are stored in the database with identifiers, timestamps, and ownership details.",
        "Verify that history or audit entries are recorded for significant user actions.",
        "Check that update and delete actions honor authorization and ownership controls.",
        "Validate that export or document-generation features produce expected outputs where supported.",
        "Review environment configuration, deployment status, and backup/restore readiness.",
        "Confirm that operational support contacts and escalation paths are documented for issue handling.",
    ]
    for idx, item in enumerate(checklist_items, start=1):
        document.add_paragraph(f"[ ] {idx}. {item}")


def add_contact_information(document: Document) -> None:
    document.add_heading("Contact Information", level=1)
    document.add_paragraph(
        "The following role-based contacts should be available to the inspection team through the internal "
        "organizational escalation matrix."
    )

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = "Contact Role"
    headers[1].text = "Primary Purpose"
    headers[2].text = "Department"
    headers[3].text = "Contact Channel"

    contacts = [
        (
            "Project Manager",
            "Project scope, timeline, and inspection coordination",
            "PMO / Business Applications",
            "Internal escalation matrix",
        ),
        (
            "Lead Developer",
            "System design, technical implementation, and defect clarification",
            "Software Engineering",
            "Internal escalation matrix",
        ),
        (
            "System Administrator",
            "Deployment, hosting, configuration, and environment issues",
            "IT Operations",
            "Internal escalation matrix",
        ),
        (
            "Operations Representative",
            "Functional walkthrough, user process, and certificate operations",
            "Certificate Processing Team",
            "Internal escalation matrix",
        ),
    ]
    for contact in contacts:
        row = table.add_row().cells
        for idx, value in enumerate(contact):
            row[idx].text = value


def main() -> None:
    document = Document()
    set_document_defaults(document)

    for section in document.sections:
        add_header_footer(section)

    add_cover_page(document)
    add_toc(document)
    add_system_overview(document)
    add_how_it_works(document)
    add_roles(document)
    add_access_levels(document)
    add_inspection_checklist(document)
    add_contact_information(document)

    document.save(OUTPUT_FILE)
    print(f"Created {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
