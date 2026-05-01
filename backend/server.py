from fastapi import FastAPI, APIRouter, HTTPException, Query, Body, Request, File, UploadFile
from starlette.middleware.cors import CORSMiddleware
from starlette.middleware.trustedhost import TrustedHostMiddleware
from starlette.middleware.httpsredirect import HTTPSRedirectMiddleware
from dotenv import load_dotenv
import os, json, logging
import ipaddress, openlocationcode
import math
import asyncio
from collections import defaultdict, deque
from contextlib import contextmanager
from pathlib import Path
from datetime import datetime, timezone, timedelta
from threading import Lock
import base64
import hashlib
import hmac
import secrets
from uuid import uuid4
from typing import Any, Dict, List, Optional, Literal, Callable, Union, Tuple
from pydantic import BaseModel, Field, ConfigDict
from fastapi.responses import Response
from sqlalchemy.orm import Session as DBSession
try:
    from .database import engine, SessionLocal, get_db
    from .models import User, Certificate, Client, History, Session as SessionModel, TemporaryAccess, OfficeLocation
except ImportError:
    from database import engine, SessionLocal, get_db
    from models import User, Certificate, Client, History, Session as SessionModel, TemporaryAccess, OfficeLocation
from sqlalchemy import text, inspect, or_, func
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

# -----------------------------
# Env + storage
# -----------------------------
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env", override=False)

PURPOSE_TEXT_FV = {
    "allotment": (
        "allotment of shares under Section 56(2)(x) of the Income Tax Act, 1961 read with "
        "Rule 11UA of the Income Tax Rules, 1962, and as required under applicable provisions "
        "of the Companies Act, 2013"
    ),
    "buyback": (
        "buy-back of shares under Section 68 of the Companies Act, 2013 and the Companies "
        "(Share Capital and Debentures) Rules, 2014"
    ),
    "merger": (
        "Merger / De-merger under Sections 230-232 of the Companies Act, 2013 and as may be "
        "required by the National Company Law Tribunal (NCLT)"
    ),
    "rtor": (
        "transfer of shares from a Resident to a Non-Resident under the Foreign Exchange "
        "Management Act, 1999 (FEMA) and the Foreign Exchange Management (Non-Debt Instruments) "
        "Rules, 2019, as per the pricing guidelines of the Reserve Bank of India"
    ),
    "form3ceb": (
        "reporting of international transactions at Arm's Length Price under Section 92 of the "
        "Income Tax Act, 1961, as required to be certified in Form 3CEB"
    ),
}

def generate_lod_docx(data: dict) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    def add_para(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                 underline=False, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold      = bold
            run.underline = underline
            run.font.size = Pt(size)
        return p

    # --- Header ---
    add_para("TO WHOM SO EVER IT MAY CONCERN", bold=True, underline=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_before=18, space_after=18)

    # Para 1
    company_name = data.get("company_name", "")
    cin = data.get("cin", "")
    address = data.get("registered_address", "")
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(12)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p1.add_run("This is to certify that the company M/s ")
    r.font.size = Pt(11)
    r_name = p1.add_run(company_name)
    r_name.font.size = Pt(11)
    r_name.bold = True
    r2 = p1.add_run(" (CIN-")
    r2.font.size = Pt(11)
    r_cin = p1.add_run(cin)
    r_cin.font.size = Pt(11)
    r_cin.bold = True
    r3 = p1.add_run(f"), registered at {address}.")
    r3.font.size = Pt(11)

    # Para 2
    as_on_date = data.get("as_on_date", "")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_p2a = p2.add_run("As per the records and information available on the website of the Ministry of Corporate Affairs (MCA) and based on the Annual Filing status of the Company, the ")
    r_p2a.font.size = Pt(11)
    r_p2b = p2.add_run("List of Director")
    r_p2b.bold = True
    r_p2b.font.size = Pt(11)
    r_p2c = p2.add_run(" of the Company as on ")
    r_p2c.font.size = Pt(11)
    r_p2d = p2.add_run(as_on_date)
    r_p2d.bold = True
    r_p2d.font.size = Pt(11)
    r_p2e = p2.add_run(" is as under:")
    r_p2e.font.size = Pt(11)

    # Table
    directors = data.get("directors", [])
    if directors:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        widths = (Inches(0.6), Inches(2.2), Inches(1.2), Inches(1.5), Inches(1.0))
        for i, width in enumerate(widths):
            table.columns[i].width = width
        
        hdr_cells = table.rows[0].cells
        headers = ["Sr.\nNo", "Name of Directors", "DIN", "Designation", "Date of\nAppointment"]
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            p.text = h
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(11)
        
        for d in directors:
            row = table.add_row().cells
            row[0].text = str(d.get("sr_no", ""))
            row[1].text = d.get("name", "")
            row[2].text = d.get("din", "")
            row[3].text = d.get("designation", "")
            row[4].text = d.get("date_of_appointment", "")
            
        doc.add_paragraph() # spacing

    # Para 3
    add_para("This certificate is issued on the basis of the records and documents produced before us and information available on the MCA portal.", size=11, space_after=12)

    # Para 4
    add_para("This certificate is issued at the specific request of the Company for submission to the Bank loan purpose only and shall not be used for any other purpose without our prior written consent.", size=11, space_after=36)

    # Footer
    add_para(f"For {data.get('caFirm', '')}", bold=True, size=11, space_after=2)
    add_para("(Chartered Accountants)", bold=True, size=11, space_after=2)
    add_para(f"F. R. No. {data.get('frn', '')}", bold=True, size=11, space_after=64)
    add_para(f"(CA {data.get('caName', '')})", bold=True, size=11, space_after=2)
    add_para("(Partner)", bold=True, size=11, space_after=2)
    add_para(f"Membership No: {data.get('membershipNo', '')}", bold=True, size=11, space_after=2)
    add_para(f"UDIN: {data.get('udin', '')}", bold=True, size=11, space_after=0)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


ENVIRONMENT = (os.getenv("ENVIRONMENT") or os.getenv("VERCEL_ENV") or "development").strip().lower()
IS_PRODUCTION = ENVIRONMENT in {"prod", "production"}

DATA_DIR = Path(os.getenv("STORAGE_DIR", str(ROOT_DIR / "data")))
if not IS_PRODUCTION:
    DATA_DIR.mkdir(parents=True, exist_ok=True)

# Office access control managed via environment and database
BACKEND_DIR = Path(__file__).parent

WEAK_DEFAULT_ADMIN_PASSWORDS = {
    "admin",
    "admin123",
    "admin@123",
    "password",
    "password123",
    "123456",
}

EntityType = Literal[
    "PERSONAL",
    "PROPRIETORSHIP",
    "PRIVATE_LIMITED",
    "PUBLIC_LIMITED",
    "TRUST",
    "NGO",
    "SOCIETY",
    "GOVERNMENT",
    "COLLEGE",
]

CertCategory = Literal[
    "NET_WORTH",
    "TURNOVER",
    "UTILISATION",
    "RERA",
    "NBFC",
    "GST",
    "LIST_OF_DIRECTORS",
]

def _now() -> datetime:
    return datetime.now(timezone.utc)

def _env_int(name: str, default: int, minimum: int = 1) -> int:
    raw = (os.getenv(name) or "").strip()
    if not raw:
        return default
    try:
        value = int(raw)
        if value < minimum:
            return default
        return value
    except ValueError:
        return default

def _env_bool(name: str, default: bool = False) -> bool:
    raw = (os.getenv(name) or "").strip().lower()
    if not raw:
        return default
    return raw in {"1", "true", "yes", "on"}

def _parse_cors_origins() -> List[str]:
    raw = os.getenv("CORS_ORIGINS", "").strip()
    if not raw:
        if IS_PRODUCTION:
            raise RuntimeError("CORS_ORIGINS is required in production.")
        return ["http://localhost:3000", "http://127.0.0.1:3000"]

    origins = [origin.strip() for origin in raw.split(",") if origin.strip()]
    if not origins:
        if IS_PRODUCTION:
            raise RuntimeError("CORS_ORIGINS is required in production.")
        return ["http://localhost:3000", "http://127.0.0.1:3000"]
    return origins

def _parse_allowed_hosts() -> List[str]:
    raw = os.getenv("ALLOWED_HOSTS", "").strip()
    if not raw:
        if IS_PRODUCTION:
            raise RuntimeError("ALLOWED_HOSTS is required in production.")
        return ["localhost", "127.0.0.1", "[::1]"]
    hosts = [host.strip() for host in raw.split(",") if host.strip()]
    if not hosts:
        if IS_PRODUCTION:
            raise RuntimeError("ALLOWED_HOSTS is required in production.")
        return ["localhost", "127.0.0.1", "[::1]"]
    return hosts

def _validate_password_strength(password: str) -> Optional[str]:
    if len(password) < 12:
        return "Password must be at least 12 characters."
    if not any(ch.isupper() for ch in password):
        return "Password must contain at least one uppercase letter."
    if not any(ch.islower() for ch in password):
        return "Password must contain at least one lowercase letter."
    if not any(ch.isdigit() for ch in password):
        return "Password must contain at least one number."
    if not any(not ch.isalnum() for ch in password):
        return "Password must contain at least one special character."
    return None

LOGIN_RATE_LIMIT_ATTEMPTS = _env_int("LOGIN_RATE_LIMIT_ATTEMPTS", 5, 1)
LOGIN_RATE_LIMIT_WINDOW_SEC = _env_int("LOGIN_RATE_LIMIT_WINDOW_SEC", 300, 1)
LOGIN_RATE_LIMIT_LOCKOUT_SEC = _env_int("LOGIN_RATE_LIMIT_LOCKOUT_SEC", 900, 1)
FORCE_HTTPS = _env_bool("FORCE_HTTPS", IS_PRODUCTION)
ENABLE_API_DOCS = _env_bool("ENABLE_API_DOCS", not IS_PRODUCTION)
_login_attempts: Dict[str, deque[datetime]] = defaultdict(deque)
_login_lockouts: Dict[str, datetime] = {}
_login_rate_limit_lock = Lock()
CORS_ALLOW_ORIGINS = _parse_cors_origins()
TRUSTED_ALLOWED_HOSTS = _parse_allowed_hosts()

# ------------------------------------------------------------------
# In-memory caches to reduce Neon round-trips
# ------------------------------------------------------------------
# Session token cache: token -> (user_dict, session_dict, cached_at)
_SESSION_CACHE_TTL_SEC = 60
_session_cache: Dict[str, Tuple[Dict[str, Any], Dict[str, Any], datetime]] = {}
_session_cache_lock = Lock()

# Office location cache: (locations_list, cached_at)
_OFFICE_CACHE_TTL_SEC = 60
_office_cache: Optional[Tuple[List[Dict[str, Any]], datetime]] = None
_office_cache_lock = Lock()

def _get_cached_session(token: str) -> Optional[Tuple[Dict[str, Any], Dict[str, Any]]]:
    """Return (user, session) from cache if still valid."""
    with _session_cache_lock:
        entry = _session_cache.get(token)
        if not entry:
            return None
        u_dict, s_dict, c_at = entry
        if (datetime.now(timezone.utc) - c_at).total_seconds() > _SESSION_CACHE_TTL_SEC:
            _session_cache.pop(token, None)
            return None
        return u_dict, s_dict

def _put_cached_session(token: str, user: Dict[str, Any], session: Dict[str, Any]) -> None:
    with _session_cache_lock:
        if len(_session_cache) > 500:
            keys = sorted(_session_cache.keys(), key=lambda k: _session_cache[k][2])
            for k in keys[:100]:
                _session_cache.pop(k, None)
        _session_cache[token] = (user, session, datetime.now(timezone.utc))

def _invalidate_cached_session(token: str) -> None:
    with _session_cache_lock:
        _session_cache.pop(token, None)

def _invalidate_cached_sessions_for_user(user_id: str) -> None:
    with _session_cache_lock:
        to_remove = [k for k, v in _session_cache.items() if v[1].get("user_id") == user_id]
        for k in to_remove:
            _session_cache.pop(k, None)

def _get_office_locations_cached() -> List[Dict[str, Any]]:
    """Return office locations from in-memory cache (max 60s stale)."""
    global _office_cache
    with _office_cache_lock:
        if _office_cache is not None:
            locations, cached_at = _office_cache
            if (datetime.now(timezone.utc) - cached_at).total_seconds() <= _OFFICE_CACHE_TTL_SEC:
                return locations
        locations = _get_office_locations()
        _office_cache = (locations, datetime.now(timezone.utc))
        return locations

def _invalidate_office_cache():
    global _office_cache
    with _office_cache_lock:
        _office_cache = None


def _validate_env():
    if not IS_PRODUCTION:
        return
    
    required = [
        "JWT_SECRET",
        "ADMIN_EMAIL",
        "ADMIN_PASSWORD",
    ]
    missing = [r for r in required if not os.getenv(r)]
    if not os.getenv("DATABASE_URL"):
        missing.append("DATABASE_URL")
    if missing:
        raise RuntimeError(f"Missing required production environment variables: {', '.join(missing)}")
    
    if IS_PRODUCTION:
        val_email = os.getenv("ADMIN_EMAIL") or os.getenv("ADMIN_USERNAME")
        val_pass = os.getenv("ADMIN_PASSWORD")
        if (not val_email or not val_pass) or (val_pass and val_pass.strip().lower() in WEAK_DEFAULT_ADMIN_PASSWORDS):
            raise RuntimeError("Weak or missing ADMIN_PASSWORD in production.")
        if val_pass:
            password_error = _validate_password_strength(val_pass)
            if password_error:
                raise RuntimeError(f"ADMIN_PASSWORD error: {password_error}")

_validate_env()



@contextmanager
def _db():
    """Transition helper: provides a session that behaves somewhat like a connection for raw SQL"""
    session = SessionLocal()
    try:
        yield session
        session.commit()
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

def _init_db() -> None:
    # Logic moved to Alembic, but we can keep create_all for safety in dev
    try:
        from .database import Base
    except ImportError:
        from database import Base
    Base.metadata.create_all(bind=engine)
    inspector = inspect(engine)

    # Development safety net: add columns introduced after initial migrations.
    # create_all() does not alter existing tables.
    if not inspector.has_table("history"):
        return

    history_columns = {col["name"] for col in inspector.get_columns("history")}
    certificates_columns = {col["name"] for col in inspector.get_columns("certificates")} if inspector.has_table("certificates") else set()

    with engine.begin() as conn:
        if "user_email" not in history_columns:
            conn.execute(text("ALTER TABLE history ADD COLUMN user_email VARCHAR"))
        if "created_by_info" not in certificates_columns:
            conn.execute(text("ALTER TABLE certificates ADD COLUMN created_by_info VARCHAR"))
        
        # SQLite doesn't support IF NOT EXISTS in CREATE INDEX for some versions 
        # or it might already exist. We try for safety.
        try:
            conn.execute(text("CREATE INDEX ix_history_user_email ON history (user_email)"))
        except Exception:
            pass



def _serialize(doc: Dict[str, Any]) -> Dict[str, Any]:
    # convert datetime to isoformat for json storage
    for k, v in list(doc.items()):
        if isinstance(v, datetime):
            doc[k] = v.isoformat()
    return doc

def _deserialize(doc: Dict[str, Any]) -> Dict[str, Any]:
    for k in ("created_at", "updated_at"):
        if isinstance(doc.get(k), str):
            try:
                doc[k] = datetime.fromisoformat(doc[k])
            except Exception:
                pass
    return doc

def _parse_dt(value: Optional[str]) -> Optional[datetime]:
    if not value or not isinstance(value, str):
        return None
    try:
        return datetime.fromisoformat(value)
    except Exception:
        return None

def _is_expired(expires_at: Optional[str]) -> bool:
    dt = _parse_dt(expires_at)
    if not dt:
        return False
    return _now() >= dt

def _hash_password(password: str) -> str:
    iterations = 120_000
    salt = os.urandom(16)
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, iterations)
    return "pbkdf2_sha256${}${}${}".format(
        iterations,
        base64.b64encode(salt).decode("utf-8"),
        base64.b64encode(dk).decode("utf-8"),
    )

def _verify_password(password: str, stored: str) -> bool:
    try:
        algo, iterations_str, salt_b64, hash_b64 = stored.split("$", 3)
        if algo != "pbkdf2_sha256":
            return False
        iterations = int(iterations_str)
        salt = base64.b64decode(salt_b64.encode("utf-8"))
        expected = base64.b64decode(hash_b64.encode("utf-8"))
        dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, iterations)
        return hmac.compare_digest(dk, expected)
    except Exception:
        return False

def _new_token() -> str:
    return secrets.token_urlsafe(32)

def _random_password() -> str:
    return secrets.token_urlsafe(12)

def _is_blank(v: Optional[str]) -> bool:
    return not (v or "").strip()


# -----------------------------
# Universal models
# -----------------------------
class UniversalIdentity(BaseModel):
    model_config = ConfigDict(extra="ignore")

    person_name: str = ""
    company_name: str = ""

    legal_type: str = ""      # optional (e.g., NGO legal type)
    reg_no: str = ""          # optional
    department: str = ""      # optional (Govt/Department)

    pan: str = ""
    cin: str = ""
    gstin: str = ""
    address: str = ""

class UniversalMeta(BaseModel):
    model_config = ConfigDict(extra="ignore")
    purpose: str = ""
    place: str = ""
    date: str = ""            # keep string (dd/mm/yyyy)

    # ✅ IMPORTANT: NetWorthForm sends this; without it, it gets dropped.
    as_on_date: str = ""      # dd/mm/yyyy

class UniversalCA(BaseModel):
    model_config = ConfigDict(extra="ignore")
    firm: str = ""
    frn: str = ""
    name: str = ""
    membership_no: str = ""
    udin: str = ""

class UniversalTable(BaseModel):
    model_config = ConfigDict(extra="ignore")

    columns: List[str] = []

    # ✅ rows can be:
    # - List[List[Any]]   → tabular data (NetWorth, Utilisation)
    # - List[Dict[str, Any]] → structured rows (RERA main_form)
    rows: List[Union[List[Any], Dict[str, Any]]] = []


class UniversalData(BaseModel):
    model_config = ConfigDict(extra="ignore")
    tables: Dict[str, UniversalTable] = {}
    extras: Dict[str, Any] = {}

class UniversalCertificateCreate(BaseModel):
    model_config = ConfigDict(extra="ignore")

    client_id: Optional[str] = None
    category: CertCategory
    certificate_type: str          # e.g. "turnover_certificate", "net_worth_certificate"
    entityType: EntityType

    identity: UniversalIdentity
    meta: UniversalMeta
    ca: UniversalCA
    data: UniversalData

class UniversalCertificateStored(BaseModel):
    model_config = ConfigDict(extra="ignore")

    id: str = Field(default_factory=lambda: str(uuid4()))
    client_id: Optional[str] = None
    category: CertCategory
    certificate_type: str
    entityType: EntityType

    identity: Dict[str, Any] = {}
    meta: Dict[str, Any] = {}
    ca: Dict[str, Any] = {}
    data: Dict[str, Any] = {}

    created_at: datetime = Field(default_factory=_now)
    updated_at: datetime = Field(default_factory=_now)

# -----------------------------
# Auth models
# -----------------------------
class AuthLoginRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    username: Optional[str] = None
    email: Optional[str] = None
    password: str
    full_name: Optional[str] = None

class AuthLoginResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    token: str
    user: Dict[str, Any]

class TempCredentialCreate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    username: Optional[str] = None
    email: Optional[str] = None
    full_name: Optional[str] = None
    password: Optional[str] = None
    expires_in_hours: Optional[int] = 12
    role: Optional[Literal["temporary", "staff", "data_executive"]] = "temporary"
    can_manage_certificates: Optional[bool] = None
    can_edit_certificates: Optional[bool] = False
    can_delete_certificates: Optional[bool] = False

class RevokeCredentialRequest(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: Optional[str] = None
    temp_access_id: Optional[str] = None

class GeoCheckRequest(BaseModel):
    lat: float
    lng: float

class OfficeLocationCreate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    name: str
    ips: Optional[List[str]] = []
    plus_code: Optional[str] = None
    lat: Optional[float] = None
    lng: Optional[float] = None
    radius_m: Optional[float] = None

class OfficeLocationUpdate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    name: Optional[str] = None
    ips: Optional[List[str]] = None
    plus_code: Optional[str] = None
    lat: Optional[float] = None
    lng: Optional[float] = None
    radius_m: Optional[float] = None

class ClientCreate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    entity_type: EntityType
    display_name: Optional[str] = None
    person_name: Optional[str] = None
    company_name: Optional[str] = None
    pan: Optional[str] = None
    cin: Optional[str] = None
    gstin: Optional[str] = None
    address: Optional[str] = None

class ClientUpdate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    entity_type: Optional[EntityType] = None
    display_name: Optional[str] = None
    person_name: Optional[str] = None
    company_name: Optional[str] = None
    pan: Optional[str] = None
    cin: Optional[str] = None
    gstin: Optional[str] = None
    address: Optional[str] = None

# -----------------------------
# Validation (common + per category)
# -----------------------------
def _required_display_name(entityType: str, identity: UniversalIdentity) -> Optional[str]:
    """
    Rule:
    - PERSONAL => person_name required
    - Others => company_name required
    """
    if entityType == "PERSONAL":
        if _is_blank(identity.person_name):
            return "person_name is required for PERSONAL."
    else:
        if _is_blank(identity.company_name):
            return "company_name is required for non-PERSONAL entities."
    return None

def validate_common(payload: UniversalCertificateCreate) -> Optional[str]:
    cert_type = payload.certificate_type.lower()

    # ---------------- ENTITY NAME ----------------
    err = _required_display_name(payload.entityType, payload.identity)
    if err:
        return err

    # ---------------- META ----------------
    # ❌ purpose NOT required for RERA Form 7
    if cert_type != "rera_form_7_reg_9":
        if _is_blank(payload.meta.purpose):
            return "purpose is required."
        
        if _is_blank(payload.meta.place):
            return "place is required."
    
        if _is_blank(payload.meta.date):
            return "date is required."

    # place & date are still required (even for RERA)


    # ---------------- CA DETAILS ----------------
    # ❌ CA details NOT required for RERA Form 7
    if cert_type != "rera_form_7_reg_9":
        if _is_blank(payload.ca.firm):
            return "CA firm is required."
        if _is_blank(payload.ca.frn):
            return "FRN is required."
        if _is_blank(payload.ca.name):
            return "CA name is required."
        if _is_blank(payload.ca.membership_no):
            return "CA membership_no is required."

    return None

def validate_turnover(payload: UniversalCertificateCreate) -> Optional[str]:
    t = payload.data.tables.get("main")
    if not t or not t.rows:
        return "TURNOVER requires data.tables.main with at least one row."
    return None

def validate_networth(payload: UniversalCertificateCreate) -> Optional[str]:
    # Net Worth requires an As On date
    if _is_blank(payload.meta.as_on_date):
        return "NET_WORTH requires meta.as_on_date."

    # Require the schedule tables coming from NetWorthForm
    required_tables = ["scheduleA", "scheduleB", "scheduleC", "summary"]
    for key in required_tables:
        t = payload.data.tables.get(key)
        if not t:
            return f"NET_WORTH requires data.tables.{key}."
        # schedules must have rows (at least 1)
        if key in ("scheduleA", "scheduleB", "scheduleC"):
            if not t.rows or len(t.rows) == 0:
                return f"NET_WORTH requires data.tables.{key}.rows with at least one row."

    return None

def validate_utilisation(payload: UniversalCertificateCreate) -> Optional[str]:
    tables = payload.data.tables or {}

    # Payment details required
    payment = tables.get("paymentDetails")
    if not payment or not payment.rows:
        return "UTILISATION requires paymentDetails table with at least one row."

    # Period required
    period = tables.get("period")
    if not period or not period.rows:
        return "UTILISATION requires period table with from and to dates."

    # Summary required
    summary = tables.get("summary")
    if not summary or not summary.rows:
        return "UTILISATION requires summary table."

    return None

def validate_rera(payload: UniversalCertificateCreate):
    cert_type = payload.certificate_type.lower()

    # ---------------- RERA FORM 3 ----------------
    if cert_type == "rera_form_3":
        tables = payload.data.tables or {}

        required_tables = ["main_form", "sold_inventory", "unsold_inventory"]
        for key in required_tables:
            t = tables.get(key)
            if not t or not isinstance(t.rows, list) or len(t.rows) == 0:
                return f"RERA Form 3 requires data.tables.{key}.rows"

        extras = payload.data.extras or {}
        if not extras.get("projectName"):
            return "RERA Form 3 requires extras.projectName"
        if not extras.get("reraRegistrationNumber"):
            return "RERA Form 3 requires extras.reraRegistrationNumber"

        return None

    # ---------------- RERA FORM 7 ----------------
    if cert_type == "rera_form_7_reg_9":
        extras = payload.data.extras or {}
        form = extras.get("formData")

        if not isinstance(form, dict):
            return "RERA Form 7 requires data.extras.formData"

        # minimal statutory checks (keep flexible)
        if not form.get("meta", {}).get("year"):
            return "RERA Form 7 requires year"
        if not form.get("projectDetails", {}).get("registrationNumber"):
            return "RERA Form 7 requires project registration number"

        return None

    return "Unknown RERA certificate type"

def validate_nbfc(payload: UniversalCertificateCreate) -> Optional[str]:
    cert_type = (payload.certificate_type or "").lower().strip()
    if cert_type != "rbi_statutory_auditor_certificate_for_nbfcs":
        return "Unknown NBFC certificate type"

    tables = payload.data.tables or {}
    main = tables.get("main")
    if not main or not main.rows:
        return "NBFC requires data.tables.main with checklist rows."

    extras = payload.data.extras or {}
    form = extras.get("formData")
    if not isinstance(form, dict):
        return "NBFC requires data.extras.formData"

    if _is_blank(str(form.get("financialYearEnd") or "")):
        return "NBFC requires financial year end."

    return None

def validate_lod(payload: UniversalCertificateCreate) -> Optional[str]:
    data = payload.data.extras or {}
    if _is_blank(payload.identity.cin):
        return "CIN is required for List of Directors."
    if _is_blank(data.get("as_on_date")):
        return "As on Date is required for List of Directors."
    if not data.get("directors") or not isinstance(data.get("directors"), list):
        return "At least one director must be listed."
    return None


CATEGORY_VALIDATORS: Dict[str, Callable[[UniversalCertificateCreate], Optional[str]]] = {
    "TURNOVER": validate_turnover,
    "NET_WORTH": validate_networth,
    "UTILISATION": validate_utilisation,
    "RERA": validate_rera,
    "NBFC": validate_nbfc,
    "LIST_OF_DIRECTORS": validate_lod,
}

# -----------------------------
# Auth helpers
# -----------------------------
def _sanitize_user(user: Dict[str, Any]) -> Dict[str, Any]:
    safe = dict(user)
    safe.pop("password_hash", None)
    if safe.get("email") and "username" not in safe:
        safe["username"] = safe["email"]
    role = (safe.get("role") or "").lower()
    if role == "admin":
        safe["can_edit_certificates"] = True
        safe["can_delete_certificates"] = True
        safe["can_manage_certificates"] = True
    else:
        safe["can_edit_certificates"] = bool(int(safe.get("can_edit_certificates") or 0))
        safe["can_delete_certificates"] = bool(int(safe.get("can_delete_certificates") or 0))
        safe["can_manage_certificates"] = bool(
            safe["can_edit_certificates"] or safe["can_delete_certificates"]
        )
    return safe

def _log_history(user_id: str, action_type: str, action_data: Optional[Dict[str, Any]] = None) -> None:
    # Try to find user email for the snapshot
    with _db() as db:
        user = db.query(User).filter(User.id == user_id).first()
        email_snapshot = user.email if user else "deleted-user"
        
        history = History(
            id=str(uuid4()),
            user_id=user_id,
            user_email=email_snapshot,
            action_type=action_type,
            action_data=json.dumps(action_data or {}, ensure_ascii=False, default=str),
            timestamp=_now().isoformat()
        )
        db.add(history)

async def _ensure_admin_user() -> None:
    admin_email = (os.getenv("ADMIN_EMAIL") or os.getenv("ADMIN_USERNAME") or "").strip().lower()
    admin_password = (os.getenv("ADMIN_PASSWORD") or "").strip()

    if not admin_email or not admin_password:
        if IS_PRODUCTION:
            raise RuntimeError("ADMIN_EMAIL/ADMIN_USERNAME and ADMIN_PASSWORD are required in production.")
        admin_email = admin_email or "admin"
        admin_password = admin_password or "admin123"

    if IS_PRODUCTION:
        if admin_password.lower() in WEAK_DEFAULT_ADMIN_PASSWORDS:
            raise RuntimeError("ADMIN_PASSWORD is too weak for production.")
        password_error = _validate_password_strength(admin_password)
        if password_error:
            raise RuntimeError(f"ADMIN_PASSWORD is invalid for production. {password_error}")

    with _db() as db:
        user = db.query(User).filter(User.email == admin_email).first()
        if user:
            user.can_edit_certificates = 1
            user.can_delete_certificates = 1
            if admin_password and not _verify_password(admin_password, user.password_hash or ""):
                user.password_hash = _hash_password(admin_password)
            db.commit()
            return

        admin_user = User(
            id=str(uuid4()),
            email=admin_email,
            full_name="Admin",
            role="admin",
            can_edit_certificates=1,
            can_delete_certificates=1,
            password_hash=_hash_password(admin_password),
            created_at=_now().isoformat()
        )
        db.add(admin_user)

def _get_user_by_email(email: str) -> Optional[Dict[str, Any]]:
    with _db() as db:
        user = db.query(User).filter(User.email == email).first()
        return _to_dict(user) if user else None

def _get_user_by_id(user_id: str) -> Optional[Dict[str, Any]]:
    with _db() as db:
        user = db.query(User).filter(User.id == user_id).first()
        return _to_dict(user) if user else None

def _to_dict(model_instance) -> Dict[str, Any]:
    """Helper to convert SA model to dict for compatibility with existing code"""
    if not model_instance:
        return {}
    d = {}
    for column in model_instance.__table__.columns:
        d[column.name] = getattr(model_instance, column.name)
    return d

def _create_user_if_missing(
    email: str,
    role: str,
    can_manage_certificates: bool = False,
) -> Dict[str, Any]:
    existing = _get_user_by_email(email)
    if existing:
        return existing
    
    user_obj = User(
        id=str(uuid4()),
        email=email,
        full_name=None,
        role=role,
        can_edit_certificates=1 if can_manage_certificates else 0,
        can_delete_certificates=1 if can_manage_certificates else 0,
        password_hash=None,
        created_at=_now().isoformat()
    )
    with _db() as db:
        db.add(user_obj)
        db.commit()
        return _to_dict(user_obj)
    return {}  # unreachable but satisfies type checker

def _update_user_full_name(user_id: str, full_name: str) -> None:
    with _db() as db:
        db.query(User).filter(User.id == user_id).update({"full_name": full_name})

def _update_user_certificate_permissions(user_id: str, can_manage: bool) -> None:
    with _db() as db:
        db.query(User).filter(User.id == user_id).update({
            "can_edit_certificates": 1 if can_manage else 0,
            "can_delete_certificates": 1 if can_manage else 0
        })

def _update_user_role(user_id: str, role: str) -> None:
    with _db() as db:
        db.query(User).filter(User.id == user_id).update({"role": role})

def _can_user_manage_certificates(user: Dict[str, Any]) -> bool:
    if (user.get("role") or "").lower() == "admin":
        return True
    return bool(
        int(user.get("can_edit_certificates") or 0)
        or int(user.get("can_delete_certificates") or 0)
    )

def _can_user_manage_clients(user: Dict[str, Any]) -> bool:
    return (user.get("role") or "").lower() in {"admin", "data_executive"}

def _clean_client_text(value: Optional[str]) -> str:
    return str(value or "").strip()

def _normalize_client_identifier(value: Optional[str]) -> str:
    return _clean_client_text(value).upper()

def _client_display_name(data: Dict[str, Any]) -> str:
    return (
        _clean_client_text(data.get("display_name"))
        or _clean_client_text(data.get("company_name"))
        or _clean_client_text(data.get("person_name"))
    )

def _serialize_client(client: Client) -> Dict[str, Any]:
    data = _to_dict(client)
    data["is_deleted"] = bool(int(data.get("is_deleted") or 0))
    return data

def _assert_no_active_client_duplicate(
    db: DBSession,
    pan: Optional[str],
    cin: Optional[str],
    gstin: Optional[str],
    exclude_id: Optional[str] = None,
) -> None:
    duplicate_checks = (
        ("pan", _normalize_client_identifier(pan), "PAN"),
        ("cin", _normalize_client_identifier(cin), "CIN"),
        ("gstin", _normalize_client_identifier(gstin), "GSTIN"),
    )
    for field_name, value, label in duplicate_checks:
        if not value:
            continue
        column = getattr(Client, field_name)
        query = db.query(Client).filter(
            Client.is_deleted == 0,
            func.upper(column) == value,
        )
        if exclude_id:
            query = query.filter(Client.id != exclude_id)
        existing = query.first()
        if existing:
            raise HTTPException(
                status_code=400,
                detail=f"An active client with this {label} already exists.",
            )

async def require_client_master_user(request: Request) -> Dict[str, Any]:
    user = await require_user(request)
    if not _can_user_manage_clients(user):
        raise HTTPException(status_code=403, detail="Client master access requires admin or data executive role.")
    return user

def _find_temp_access_for_login(user_id: str, password: str) -> Tuple[Optional[Dict[str, Any]], Optional[str]]:
    with _db() as db:
        access_list = db.query(TemporaryAccess).filter(TemporaryAccess.user_id == user_id).order_by(TemporaryAccess.created_at.desc()).all()
        if not access_list:
            return None, "no_access"
        for access in access_list:
            if _verify_password(password, access.password_hash):
                if access.is_revoked:
                    return None, "revoked"
                if _is_expired(access.expires_at):
                    return None, "expired"
                return _to_dict(access), None
    return None, "invalid"

def _create_temp_access(user_id: str, password_hash: str, expires_at: str, admin_id: str) -> Dict[str, Any]:
    access = TemporaryAccess(
        id=str(uuid4()),
        user_id=user_id,
        password_hash=password_hash,
        expires_at=expires_at,
        is_revoked=0,
        created_by_admin_id=admin_id,
        created_at=_now().isoformat()
    )
    with _db() as db:
        db.add(access)
        db.commit()
        return _to_dict(access)
    return {}  # unreachable

def _create_session(user_id: str, temp_access_id: Optional[str], expires_at: datetime) -> Dict[str, Any]:
    sess_obj = SessionModel(
        id=str(uuid4()),
        user_id=user_id,
        temp_access_id=temp_access_id,
        token=_new_token(),
        expires_at=expires_at.isoformat(),
        is_revoked=0,
        geo_granted_until=None,
        created_at=_now().isoformat()
    )
    with _db() as db:
        db.add(sess_obj)
        db.commit()
        return _to_dict(sess_obj)
    return {}  # unreachable

def _get_session_by_token(token: str) -> Optional[Dict[str, Any]]:
    with _db() as db:
        sess = db.query(SessionModel).filter(SessionModel.token == token).first()
        return _to_dict(sess) if sess else None

def _get_temp_access_by_id(access_id: str) -> Optional[Dict[str, Any]]:
    with _db() as db:
        access = db.query(TemporaryAccess).filter(TemporaryAccess.id == access_id).first()
        return _to_dict(access) if access else None

def _revoke_sessions_for_user(user_id: str) -> None:
    _invalidate_cached_sessions_for_user(user_id)
    with _db() as db:
        db.query(SessionModel).filter(SessionModel.user_id == user_id).update({"is_revoked": 1})

def _revoke_sessions_for_access(access_id: str) -> None:
    with _db() as db:
        db.query(SessionModel).filter(SessionModel.temp_access_id == access_id).update({"is_revoked": 1})

def _prune_expired_sessions() -> None:
    with _db() as db:
        db.query(SessionModel).filter(SessionModel.expires_at <= _now().isoformat()).delete()

def _get_client_ip(request: Request) -> Optional[str]:
    # Prefer proxy headers commonly set on shared hosting/reverse proxies
    forwarded = request.headers.get("x-forwarded-for", "")
    if forwarded:
        return forwarded.split(",")[0].strip() or None
    real_ip = request.headers.get("x-real-ip", "").strip()
    if real_ip:
        return real_ip
    if request.client:
        return request.client.host
    return None

def _login_rate_limit_key(email: str, request: Request) -> str:
    ip = _get_client_ip(request) or "unknown"
    return f"{email}|{ip}"

def _trim_login_attempts(attempts: deque[datetime], now: datetime) -> None:
    window_start = now - timedelta(seconds=LOGIN_RATE_LIMIT_WINDOW_SEC)
    while attempts and attempts[0] < window_start:
        attempts.popleft()

def _enforce_login_rate_limit(key: str) -> None:
    now = _now()
    with _login_rate_limit_lock:
        lock_until = _login_lockouts.get(key)
        if lock_until and lock_until > now:
            retry_after = max(1, int((lock_until - now).total_seconds()))
            raise HTTPException(
                status_code=429,
                detail=f"Too many login attempts. Try again in {retry_after} seconds.",
            )
        if lock_until and lock_until <= now:
            _login_lockouts.pop(key, None)

        attempts = _login_attempts.get(key, deque[datetime]())
        _login_attempts[key] = attempts # Ensure deque exists for key
        _trim_login_attempts(attempts, now)
        if len(attempts) >= LOGIN_RATE_LIMIT_ATTEMPTS:
            lock_until = now + timedelta(seconds=LOGIN_RATE_LIMIT_LOCKOUT_SEC)
            _login_lockouts[key] = lock_until
            attempts.clear()
            raise HTTPException(
                status_code=429,
                detail="Too many login attempts. Please try again later.",
            )

def _record_failed_login(key: str) -> None:
    now = _now()
    with _login_rate_limit_lock:
        attempts = _login_attempts.get(key, deque[datetime]())
        _login_attempts[key] = attempts # Ensure deque exists for key
        _trim_login_attempts(attempts, now)
        attempts.append(now)
        if len(attempts) >= LOGIN_RATE_LIMIT_ATTEMPTS:
            _login_lockouts[key] = now + timedelta(seconds=LOGIN_RATE_LIMIT_LOCKOUT_SEC)
            attempts.clear()

def _clear_login_rate_limit(key: str) -> None:
    with _login_rate_limit_lock:
        _login_attempts.pop(key, None)
        _login_lockouts.pop(key, None)

def _parse_ip_list(values: Optional[List[str]]) -> List[ipaddress._BaseNetwork]:
    if not values:
        return []
    networks: List[ipaddress._BaseNetwork] = []
    for item in values:
        if not item:
            continue
        try:
            networks.append(ipaddress.ip_network(item.strip(), strict=False))
        except ValueError:
            continue
    return networks

def _parse_ips_from_csv(value: Optional[str]) -> List[str]:
    if not value:
        return []
    return [v.strip() for v in value.split(",") if v.strip()]

def _extract_plus_code(value: Optional[str]) -> str:
    """
    Accept values like:
    - 'J47V+HC Patna, Bihar'
    - '7J4VJ47V+HC'
    and return the code token.
    """
    raw = " ".join(str(value or "").strip().upper().split())
    if not raw:
        return ""
    token = raw.split(" ", 1)[0]
    return token.strip(",;")

def _parse_lat_lng_pair(lat_raw: Optional[Any], lng_raw: Optional[Any]) -> Optional[Tuple[float, float]]:
    if not lat_raw or not lng_raw:
        return None
    try:
        return float(str(lat_raw)), float(str(lng_raw))
    except (ValueError, TypeError):
        return None

def _default_plus_code_reference() -> Optional[Tuple[float, float]]:
    # First fallback: configured env office coordinates (if present)
    env_pair = _parse_lat_lng_pair(os.getenv("OFFICE_LAT"), os.getenv("OFFICE_LNG"))
    if env_pair:
        return env_pair

    # Second fallback: first saved office with geo
    for _, lat, lng, _ in _get_office_geos():
        return lat, lng

def _decode_plus_code_to_lat_lng(
    plus_code_input: str,
    ref_lat: Optional[float] = None,
    ref_lng: Optional[float] = None,
) -> Tuple[float, float]:
    try:
        from openlocationcode import openlocationcode as olc
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Plus Code decoding is unavailable. Install dependency: openlocationcode.",
        )

    code = _extract_plus_code(plus_code_input)
    if not code or "+" not in code:
        raise HTTPException(status_code=400, detail="Invalid Plus Code format.")

    if olc.isFull(code):
        full_code = code
    elif olc.isShort(code):
        if ref_lat is None or ref_lng is None:
            ref = _default_plus_code_reference()
            if not ref:
                raise HTTPException(
                    status_code=400,
                    detail="Short Plus Code needs a nearby reference. Provide lat/lng once or set OFFICE_LAT/OFFICE_LNG.",
                )
            ref_lat, ref_lng = ref
        if ref_lat is not None and ref_lng is not None:
            full_code = olc.recoverNearest(code, float(ref_lat), float(ref_lng))
        else:
            # Fallback path if somehow None reached here
            full_code = code
    else:
        raise HTTPException(status_code=400, detail="Invalid Plus Code.")

    area = olc.decode(full_code)
    lat = (area.latitudeLo + area.latitudeHi) / 2.0
    lng = (area.longitudeLo + area.longitudeHi) / 2.0
    return lat, lng

def _get_office_locations() -> List[Dict[str, Any]]:
    with _db() as db:
        offices = db.query(OfficeLocation).order_by(OfficeLocation.created_at.asc()).all()
        items = [_to_dict(o) for o in offices]
    for item in items:
        item["ips"] = _parse_ips_from_csv(item.get("ips"))
    return items

def _get_office_networks() -> List[ipaddress._BaseNetwork]:
    locations = _get_office_locations_cached()
    if locations:
        ips = []
        for loc in locations:
            ips.extend(loc.get("ips") or [])
        return _parse_ip_list(ips)

    # Fallback to env if no DB locations exist
    raw = (os.getenv("OFFICE_IPS") or "").strip()
    return _parse_ip_list(_parse_ips_from_csv(raw))

def _is_office_ip(ip_value: Optional[str]) -> bool:
    if not ip_value:
        return False
    try:
        ip = ipaddress.ip_address(ip_value)
    except ValueError:
        return False
    for net in _get_office_networks():
        if ip in net:
            return True
    return False

def _get_office_geos() -> List[Tuple[str, float, float, float]]:
    locations = _get_office_locations_cached()
    geos: List[Tuple[str, float, float, float]] = []
    for loc in locations:
        if loc.get("lat") is None or loc.get("lng") is None:
            continue
        radius = loc.get("radius_m")
        if radius is None:
            continue
        geos.append((loc.get("name") or "Office", float(loc.get("lat") or 0), float(loc.get("lng") or 0), float(radius)))
    if geos:
        return geos

    lat_raw = os.getenv("OFFICE_LAT")
    lng_raw = os.getenv("OFFICE_LNG")
    radius_raw = os.getenv("OFFICE_RADIUS_M", "300")
    if not lat_raw or not lng_raw:
        return []
    try:
        lat = float(str(lat_raw))
        lng = float(str(lng_raw))
        radius = float(str(radius_raw))
    except (ValueError, TypeError):
        return []
    return [("Office", lat, lng, radius)]

def _haversine_meters(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    # Great-circle distance between two points in meters
    r = 6371000.0
    phi1 = math.radians(lat1)
    phi2 = math.radians(lat2)
    dphi = math.radians(lat2 - lat1)
    dlambda = math.radians(lon2 - lon1)
    a = math.sin(dphi / 2) ** 2 + math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2) ** 2
    return 2 * r * math.atan2(math.sqrt(a), math.sqrt(1 - a))

def _is_within_office_geo(lat: float, lng: float) -> Tuple[bool, Optional[float], Optional[str]]:
    offices = _get_office_geos()
    if not offices:
        return (False, None, None)
    closest_distance: float = float('inf')
    closest_name: Optional[str] = None
    for name, office_lat, office_lng, radius in offices:
        distance = _haversine_meters(lat, lng, office_lat, office_lng)
        if distance <= radius:
            return (True, distance, name)
        if distance < closest_distance:
            closest_distance = distance
            closest_name = name
    return (False, None if math.isinf(closest_distance) else closest_distance, closest_name)

def _is_staff_role(role: Optional[str]) -> bool:
    return (role or "").lower() in ("temporary", "staff", "temp", "data_executive")

def _session_geo_granted(session: Dict[str, Any]) -> bool:
    value = session.get("geo_granted_until")
    if not value:
        return False
    return not _is_expired(value)

def _grant_geo_for_session(session_id: str) -> Dict[str, Any]:
    minutes = int(os.getenv("GEO_GRANT_MINUTES", "30"))
    expires_at = _now() + timedelta(minutes=minutes)
    with _db() as db:
        db.query(SessionModel).filter(SessionModel.id == session_id).update({
            "geo_granted_until": expires_at.isoformat()
        })
    return {"geo_granted_until": expires_at.isoformat(), "minutes": minutes}

async def _auth_from_request(request: Request) -> Tuple[Optional[Dict[str, Any]], Optional[Dict[str, Any]]]:
    auth = request.headers.get("authorization", "")
    if not auth.lower().startswith("bearer "):
        return (None, None)
    token = auth.split(" ", 1)[1].strip()
    if not token:
        return (None, None)

    # --- Fast path: check in-memory cache first (avoids 3 DB round-trips) ---
    cached = _get_cached_session(token)
    if cached is not None:
        cached_user, cached_session = cached
        if not cached_session.get("is_revoked") and not _is_expired(cached_session.get("expires_at")):
            return (cached_user, cached_session)
        _invalidate_cached_session(token)

    # --- Slow path: DB lookup ---
    session = _get_session_by_token(token)
    if not session:
        return (None, None)
    if session.get("is_revoked") or _is_expired(session.get("expires_at")):
        return (None, None)

    user_id = session.get("user_id")
    if not user_id or not isinstance(user_id, str):
        return (None, None)

    user = _get_user_by_id(user_id)
    if not user:
        return (None, None)

    temp_access_id = session.get("temp_access_id")
    if temp_access_id:
        access = _get_temp_access_by_id(temp_access_id)
        if not access or access.get("is_revoked") or _is_expired(access.get("expires_at")):
            return (None, None)

    # Store in cache for subsequent requests
    _put_cached_session(token, user, session)
    return (user, session)

async def require_user_raw(request: Request) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    user, session = await _auth_from_request(request)
    if user is not None and session is not None:
        return user, session
    raise HTTPException(status_code=401, detail="Unauthorized or expired credentials.")

async def require_user(request: Request) -> Dict[str, Any]:
    user, session = await require_user_raw(request)
    # Admin bypasses all access checks
    if user.get("role") == "admin":
        return user

    if _is_staff_role(user.get("role")):
        client_ip = _get_client_ip(request)
        if _is_office_ip(client_ip) or _session_geo_granted(session):
            return user
        raise HTTPException(
            status_code=403,
            detail="Access denied. You must be on office Wi-Fi or within the allowed geo-radius.",
        )

    # Unknown roles default to denied
    raise HTTPException(status_code=403, detail="Access denied for this role.")

async def require_admin(request: Request) -> Dict[str, Any]:
    user = await require_user(request)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required.")
    return user

# -----------------------------
# App + Router
# -----------------------------
app = FastAPI(
    docs_url="/docs" if ENABLE_API_DOCS else None,
    redoc_url="/redoc" if ENABLE_API_DOCS else None,
    openapi_url="/openapi.json" if ENABLE_API_DOCS else None,
    redirect_slashes=False,
)
api_router = APIRouter(prefix="/api")

@app.on_event("startup")
async def startup_event():
    _init_db()
    await _ensure_admin_user()
    # Warmup caches and prune sessions in background, not blocking startup
    asyncio.create_task(asyncio.to_thread(_get_office_locations_cached))
    asyncio.create_task(asyncio.to_thread(_prune_expired_sessions))

@app.middleware("http")
async def auth_expiry_middleware(request: Request, call_next):
    auth = request.headers.get("authorization", "")
    if auth.lower().startswith("bearer "):
        user, session = await _auth_from_request(request)
        if not user or not session:
            return Response(
                content=json.dumps({"detail": "Unauthorized or expired credentials."}),
                status_code=401,
                media_type="application/json",
            )
        if user is not None and session is not None:
            request.state.user_id = user.get("id")
            request.state.session_id = session.get("id")
    return await call_next(request)

@app.middleware("http")
async def security_headers_middleware(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=(self)"
    response.headers["Cross-Origin-Resource-Policy"] = "same-site"
    response.headers["Cross-Origin-Opener-Policy"] = "same-origin"
    response.headers["Cache-Control"] = "no-store"
    if IS_PRODUCTION:
        forwarded_proto = (request.headers.get("x-forwarded-proto") or "").lower()
        is_https = request.url.scheme == "https" or forwarded_proto == "https"
        if is_https:
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    return response

@api_router.get("/")
async def root():
    return {"message": "Universal Certificate API"}

@api_router.get("/health")
async def health_check():
    health = {"status": "healthy", "timestamp": _now().isoformat()}
    try:
        with _db() as db:
            db.execute(text("SELECT 1"))
        health["database"] = "connected"
    except Exception as e:
        health["status"] = "unhealthy"
        health["database"] = f"error: {str(e)}"
    
    if health["status"] == "unhealthy":
        raise HTTPException(status_code=503, detail=health)
    return health



@api_router.get("/certificates/generate-docx/{cert_id}")
async def generate_cert_docx(cert_id: str, request: Request):
    user = await require_user(request)
    with _db() as db:
        cert_row = db.query(Certificate).filter(Certificate.id == cert_id).first()
        if not cert_row:
            raise HTTPException(status_code=404, detail="Certificate not found")
        
        if user.get("role") != "admin" and cert_row.user_id != user["id"]:
            raise HTTPException(status_code=403, detail="Forbidden")
        
        cert_data = json.loads(cert_row.payload_json)
        category = cert_data.get("category")
        
        if category in ("FAIR_VALUE_SHARES", "LIST_OF_DIRECTORS"):
            # Flatten data for generator
            gen_data = {
                **cert_data.get("identity", {}),
                **cert_data.get("meta", {}),
                **cert_data.get("ca", {}),
                **cert_data.get("data", {}).get("extras", {}),
                "directors": cert_data.get("data", {}).get("extras", {}).get("directors", []),
                "certificate_date": cert_data.get("meta", {}).get("date"),
                "certificate_place": cert_data.get("meta", {}).get("place"),
                "firm_name": cert_data.get("ca", {}).get("firm"),
                "firm_frn": cert_data.get("ca", {}).get("frn"),
                "ca_name": cert_data.get("ca", {}).get("name"),
                "ca_membership_no": cert_data.get("ca", {}).get("membership_no"),
            }
            docx_bytes = generate_lod_docx(gen_data)
            
            filename = f"LOD_Certificate_{cert_id}.docx"
            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        else:
            raise HTTPException(status_code=400, detail="DOCX generation not implemented for this category yet.")

@api_router.get("/ca-settings")
async def get_ca_settings(request: Request):
    await require_user(request)
    return {
        "place": os.getenv("CA_PLACE", ""),
        "firm_name": os.getenv("CA_FIRM_NAME", ""),
        "frn": os.getenv("CA_FRN", ""),
        "ca_name": os.getenv("CA_NAME", ""),
        "membership_no": os.getenv("CA_MEMBERSHIP_NO", ""),
        "udin": os.getenv("CA_UDIN", ""),
        "cas": [],
        "default_ca_id": "",
    }

# -----------------------------
# Auth endpoints
# -----------------------------
@api_router.post("/auth/login", response_model=AuthLoginResponse)
@api_router.post("/auth/login/", response_model=AuthLoginResponse)
async def auth_login(payload: AuthLoginRequest, request: Request):
    email = (payload.email or payload.username or "").strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="Email is required.")

    rate_limit_key = _login_rate_limit_key(email, request)
    _enforce_login_rate_limit(rate_limit_key)

    try:
        user = _get_user_by_email(email)
        if not user:
            raise HTTPException(status_code=401, detail="Invalid email or password.")

        if user.get("role") == "admin":
            if not await asyncio.to_thread(_verify_password, payload.password, user.get("password_hash") or ""):
                raise HTTPException(status_code=401, detail="Invalid email or password.")

            token_hours = int(os.getenv("TOKEN_TTL_HOURS", "12"))
            token_expires = _now() + timedelta(hours=token_hours)
            # Pruned in background
            session = _create_session(user["id"], None, token_expires)
            _clear_login_rate_limit(rate_limit_key)
            _log_history(user["id"], "LOGIN", {"email": email, "role": "admin"})
            resp_dict = {"token": str(session.get("token") or ""), "user": _sanitize_user(user)}
            return AuthLoginResponse.model_validate(resp_dict)

        # Temporary user login: must match an active temp access entry
        access, reason = await asyncio.to_thread(_find_temp_access_for_login, user["id"], payload.password)
        if not access:
            if reason == "revoked":
                raise HTTPException(status_code=401, detail="Credentials have been revoked.")
            if reason == "expired":
                raise HTTPException(status_code=401, detail="Credentials have expired.")
            raise HTTPException(status_code=401, detail="Invalid email or password.")

        # If full name is still missing but provided during login, store it (optional)
        full_name_provided = payload.full_name
        if not user.get("full_name") and full_name_provided and not _is_blank(full_name_provided):
            _update_user_full_name(user["id"], full_name_provided.strip())
            refreshed = _get_user_by_id(user["id"])
            if refreshed:
                user = refreshed

        token_hours = int(os.getenv("TOKEN_TTL_HOURS", "12"))
        token_expires = _now() + timedelta(hours=token_hours)
        access_expires = _parse_dt(access.get("expires_at"))
        if access_expires and access_expires < token_expires:
            token_expires = access_expires

        # Pruned in background
        session = _create_session(user["id"], access["id"], token_expires)
        _clear_login_rate_limit(rate_limit_key)
        _log_history(user["id"], "LOGIN", {"email": email, "role": user.get("role") or "temporary"})
        resp_dict = {"token": str(session.get("token") or ""), "user": _sanitize_user(user)}
        return AuthLoginResponse.model_validate(resp_dict)
    except HTTPException as exc:
        if exc.status_code == 401:
            _record_failed_login(rate_limit_key)
        raise

@api_router.post("/auth/temp-credentials")
async def create_temp_credential(payload: TempCredentialCreate, request: Request):
    admin = await require_admin(request)

    email = (payload.email or payload.username or "").strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="email is required.")

    hours_val = 12
    p_hours = payload.expires_in_hours
    if p_hours is not None:
        hours_val = p_hours
    if hours_val != 0 and (hours_val < 1 or hours_val > 168):
        raise HTTPException(status_code=400, detail="expires_in_hours must be 0 (permanent) or between 1 and 168.")

    existing = _get_user_by_email(email)
    if existing and existing.get("role") == "admin":
        raise HTTPException(status_code=400, detail="Cannot create temporary access for admin email.")

    requested_role = payload.role or "temporary"
    if requested_role not in ("temporary", "staff", "data_executive"):
        raise HTTPException(status_code=400, detail="role must be temporary, staff, or data_executive.")

    if payload.can_manage_certificates is None:
        can_manage_certificates = bool(payload.can_edit_certificates) or bool(
            payload.can_delete_certificates
        )
    else:
        can_manage_certificates = bool(payload.can_manage_certificates)

    if existing:
        user = existing
        if (user.get("role") or "").lower() == "admin":
            raise HTTPException(status_code=400, detail="Cannot modify admin permissions here.")
        if (user.get("role") or "").lower() != requested_role:
            _update_user_role(user["id"], requested_role)
        _update_user_certificate_permissions(user["id"], can_manage_certificates)
        user = _get_user_by_id(user["id"])
    else:
        user = _create_user_if_missing(
            email,
            requested_role,
            can_manage_certificates=can_manage_certificates,
        )
    if user is not None:
        full_name = payload.full_name
        if full_name and not _is_blank(full_name) and not user.get("full_name"):
            _update_user_full_name(user["id"], full_name.strip())
            user = _get_user_by_id(user["id"])

    if payload.password is not None and _is_blank(payload.password):
        raise HTTPException(status_code=400, detail="password cannot be blank.")
    payload_password = payload.password
    if payload_password:
        password_error = _validate_password_strength(str(payload_password))
        if password_error:
            raise HTTPException(status_code=400, detail=password_error)

    password = payload_password or _random_password()
    if hours_val == 0:
        expires_at = (_now() + timedelta(days=3650)).isoformat()
    else:
        expires_at = (_now() + timedelta(hours=hours_val)).isoformat()

    if user is None:
        raise HTTPException(status_code=500, detail="User could not be resolved.")
    
    admin_id = str(admin.get("id") or "")
    hashed = await asyncio.to_thread(_hash_password, password)
    access = _create_temp_access(user["id"], hashed, expires_at, admin_id)
    _log_history(user["id"], "TEMP_ACCESS_CREATED", {"email": email, "expires_at": expires_at})
    _log_history(admin["id"], "ADMIN_CREATED_TEMP_ACCESS", {"email": email, "access_id": access["id"]})

    assert user is not None
    return {
        "user": _sanitize_user(user),
        "temporary_password": password,
        "expires_at": access["expires_at"],
    }

@api_router.post("/auth/revoke")
async def revoke_credential(payload: RevokeCredentialRequest, request: Request):
    admin = await require_admin(request)
    tid = payload.temp_access_id
    if tid:
        access = _get_temp_access_by_id(tid)
        if not access:
            raise HTTPException(status_code=404, detail="Temporary access not found.")
        target_uid = str(access.get("user_id") or "")
        if not target_uid:
             raise HTTPException(status_code=404, detail="Temporary access linked user not found.")
             
        with _db() as db:
            db.query(TemporaryAccess).filter(TemporaryAccess.user_id == target_uid).update({"is_revoked": 1})
            db.commit()
        _revoke_sessions_for_user(target_uid)
        _log_history(target_uid, "TEMP_ACCESS_REVOKED", {"access_id": tid})
        _log_history(admin["id"], "ADMIN_REVOKED_TEMP_ACCESS", {"access_id": tid})
        return {"ok": True, "temp_access_id": tid}

    uid = payload.user_id
    if uid:
        user = _get_user_by_id(uid)
        if not user:
            raise HTTPException(status_code=404, detail="User not found.")
        if user.get("role") == "admin":
            raise HTTPException(status_code=400, detail="Cannot revoke admin account.")
        with _db() as db:
            db.query(TemporaryAccess).filter(TemporaryAccess.user_id == uid).update({"is_revoked": 1})
            db.commit()
        _revoke_sessions_for_user(uid)
        _log_history(uid, "TEMP_ACCESS_REVOKED", {"user_id": uid})
        _log_history(admin["id"], "ADMIN_REVOKED_TEMP_ACCESS", {"user_id": uid})
        return {"ok": True, "user_id": uid}

    raise HTTPException(status_code=400, detail="user_id or temp_access_id is required.")

@api_router.get("/admin/offices")
async def list_offices(request: Request):
    await require_admin(request)
    return {"items": _get_office_locations()}

@api_router.post("/admin/offices")
async def create_office(payload: OfficeLocationCreate, request: Request):
    await require_admin(request)
    if _is_blank(payload.name):
        raise HTTPException(status_code=400, detail="name is required.")

    lat = payload.lat
    lng = payload.lng
    if not _is_blank(payload.plus_code):
        lat, lng = _decode_plus_code_to_lat_lng(payload.plus_code or "", payload.lat, payload.lng)

    now = _now().isoformat()
    office_id = str(uuid4())
    ips_csv = ",".join([i.strip() for i in (payload.ips or []) if i and i.strip()])
    
    office = OfficeLocation(
        id=office_id,
        name=payload.name.strip(),
        ips=ips_csv,
        lat=lat,
        lng=lng,
        radius_m=payload.radius_m,
        created_at=now,
        updated_at=now
    )
    with _db() as db:
        db.add(office)
        db.commit()
    return {"id": office_id}

@api_router.put("/admin/offices/{office_id}")
async def update_office(office_id: str, payload: OfficeLocationUpdate, request: Request):
    await require_admin(request)
    with _db() as db:
        office = db.query(OfficeLocation).filter(OfficeLocation.id == office_id).first()
        if not office:
            raise HTTPException(status_code=404, detail="Office location not found.")

        if payload.name is not None:
            name_val = str(payload.name).strip()
            if _is_blank(name_val):
                raise HTTPException(status_code=400, detail="name is required.")
            office.name = name_val
            
        if payload.ips is not None:
            ips_list = payload.ips or []
            office.ips = ",".join([str(i).strip() for i in ips_list if i and str(i).strip()])

        lat = payload.lat if payload.lat is not None else office.lat
        lng = payload.lng if payload.lng is not None else office.lng
        if not _is_blank(payload.plus_code):
            lat, lng = _decode_plus_code_to_lat_lng(payload.plus_code or "", lat, lng)
        
        office.lat = lat
        office.lng = lng

        if payload.radius_m is not None:
            office.radius_m = payload.radius_m
            
        office.updated_at = _now().isoformat()
        db.commit()
        
    return {"ok": True, "id": office_id}

@api_router.delete("/admin/offices/{office_id}")
async def delete_office(office_id: str, request: Request):
    await require_admin(request)
    with _db() as db:
        office = db.query(OfficeLocation).filter(OfficeLocation.id == office_id).first()
        if not office:
            raise HTTPException(status_code=404, detail="Office location not found.")
        db.delete(office)
        db.commit()
    return {"ok": True, "id": office_id}

@api_router.get("/access/status")
async def access_status(request: Request):
    user, session = await require_user_raw(request)
    role = (user.get("role") or "").lower()

    if role == "admin":
        return {"allowed": True, "role": role, "method": "admin"}

    if _is_staff_role(role):
        client_ip = _get_client_ip(request)
        if _is_office_ip(client_ip):
            return {"allowed": True, "role": role, "method": "ip"}
        if _session_geo_granted(session):
            return {"allowed": True, "role": role, "method": "geo_grant"}
        if not _get_office_geos():
            return {
                "allowed": False,
                "role": role,
                "geo_required": False,
                "message": "No office locations are configured. Please contact admin.",
            }
        return {
            "allowed": False,
            "role": role,
            "geo_required": True,
            "message": "Office Wi-Fi not detected. Please allow location to continue.",
        }

    return {"allowed": False, "role": role, "message": "Access denied for this role."}

@api_router.post("/access/geo-check")
async def access_geo_check(payload: GeoCheckRequest, request: Request):
    user, session = await require_user_raw(request)
    role = (user.get("role") or "").lower()
    if role == "admin":
        return {"allowed": True, "role": role, "method": "admin"}
    if not _is_staff_role(role):
        raise HTTPException(status_code=403, detail="Access denied for this role.")

    office_list = _get_office_geos()
    if not office_list:
        raise HTTPException(status_code=500, detail="Office geo configuration is missing.")

    within, distance, office_name = _is_within_office_geo(payload.lat, payload.lng)
    if not within:
        msg = "Access denied. You are outside the allowed geo-radius."
        if distance is not None:
            msg = f"{msg} Distance: {int(distance)} meters."
        raise HTTPException(status_code=403, detail=msg)

    grant = _grant_geo_for_session(session["id"])
    return {
        "allowed": True,
        "role": role,
        "method": "geo",
        "office": office_name,
        "geo_granted_until": grant["geo_granted_until"],
    }

@api_router.get("/auth/temp-credentials")
async def list_temp_credentials(
    request: Request,
    status: Optional[str] = Query(default="all"),
):
    await require_admin(request)
    items = []
    with _db() as db:
        results = (
            db.query(TemporaryAccess, User.email, User.role)
            .join(User, User.id == TemporaryAccess.user_id)
            .order_by(TemporaryAccess.created_at.desc())
            .all()
        )
        for ta_obj, email, role in results:
            ta_dict = _to_dict(ta_obj)
            is_expired = _is_expired(ta_dict["expires_at"])
            item = {
                "id": ta_dict["id"],
                "user_id": ta_dict["user_id"],
                "username": email,
                "email": email,
                "role": role,
                "expires_at": ta_dict["expires_at"],
                "is_revoked": bool(ta_dict["is_revoked"]),
                "is_expired": is_expired,
                "is_active": (not ta_dict["is_revoked"]) and (not is_expired),
                "created_at": ta_dict["created_at"],
            }
            items.append(item)

    if status == "active":
        items = [i for i in items if i["is_active"]]
    elif status == "expired":
        items = [i for i in items if i["is_expired"]]

    return {"items": items, "status": status}

@api_router.get("/auth/users")
async def list_users(request: Request):
    await require_admin(request)
    with _db() as db:
        users = db.query(User).order_by(User.created_at.desc()).all()
        items = [_sanitize_user(_to_dict(u)) for u in users]
    for item in items:
        item["username"] = item.get("email")
    return {"items": items}

@api_router.get("/clients")
async def list_clients(
    request: Request,
    q: Optional[str] = Query(default=None),
    entity_type: Optional[EntityType] = Query(default=None),
    limit: int = Query(default=50, ge=1, le=100),
):
    await require_client_master_user(request)
    query_text = _clean_client_text(q)
    with _db() as db:
        query = db.query(Client).filter(Client.is_deleted == 0)
        if entity_type:
            query = query.filter(Client.entity_type == entity_type)
        if query_text:
            pattern = f"%{query_text}%"
            query = query.filter(
                or_(
                    Client.display_name.ilike(pattern),
                    Client.person_name.ilike(pattern),
                    Client.company_name.ilike(pattern),
                    Client.pan.ilike(pattern),
                    Client.cin.ilike(pattern),
                    Client.gstin.ilike(pattern),
                )
            )
        clients = query.order_by(Client.display_name.asc()).limit(limit).all()
        return {"items": [_serialize_client(client) for client in clients]}

@api_router.get("/clients/{client_id}")
async def get_client(client_id: str, request: Request):
    await require_client_master_user(request)
    with _db() as db:
        client = db.query(Client).filter(Client.id == client_id, Client.is_deleted == 0).first()
        if not client:
            raise HTTPException(status_code=404, detail="Client not found.")
        return _serialize_client(client)

@api_router.post("/clients")
async def create_client(payload: ClientCreate, request: Request):
    user = await require_client_master_user(request)
    data = payload.model_dump()
    clean_data = {
        "entity_type": data["entity_type"],
        "display_name": _clean_client_text(data.get("display_name")),
        "person_name": _clean_client_text(data.get("person_name")),
        "company_name": _clean_client_text(data.get("company_name")),
        "pan": _normalize_client_identifier(data.get("pan")),
        "cin": _normalize_client_identifier(data.get("cin")),
        "gstin": _normalize_client_identifier(data.get("gstin")),
        "address": _clean_client_text(data.get("address")),
    }
    clean_data["display_name"] = _client_display_name(clean_data)
    if not clean_data["display_name"]:
        raise HTTPException(status_code=400, detail="Client name is required.")

    now = _now().isoformat()
    client = Client(
        id=str(uuid4()),
        **clean_data,
        created_by=user["id"],
        updated_by=user["id"],
        is_deleted=0,
        created_at=now,
        updated_at=now,
    )
    with _db() as db:
        _assert_no_active_client_duplicate(db, client.pan, client.cin, client.gstin)
        db.add(client)
        db.commit()
        db.refresh(client)
        result = _serialize_client(client)

    _log_history(user["id"], "CLIENT_CREATE", {"client_id": result["id"], "display_name": result["display_name"]})
    return result

@api_router.put("/clients/{client_id}")
async def update_client(client_id: str, payload: ClientUpdate, request: Request):
    user = await require_client_master_user(request)
    updates = payload.model_dump(exclude_unset=True)
    with _db() as db:
        client = db.query(Client).filter(Client.id == client_id, Client.is_deleted == 0).first()
        if not client:
            raise HTTPException(status_code=404, detail="Client not found.")

        if "entity_type" in updates:
            if _is_blank(updates.get("entity_type")):
                raise HTTPException(status_code=400, detail="entity_type is required.")
            client.entity_type = _clean_client_text(updates.get("entity_type"))
        for key in ("display_name", "person_name", "company_name", "address"):
            if key in updates:
                setattr(client, key, _clean_client_text(updates.get(key)))
        for key in ("pan", "cin", "gstin"):
            if key in updates:
                setattr(client, key, _normalize_client_identifier(updates.get(key)))

        effective_data = {
            "display_name": client.display_name,
            "company_name": client.company_name,
            "person_name": client.person_name,
        }
        client.display_name = _client_display_name(effective_data)
        if not client.display_name:
            raise HTTPException(status_code=400, detail="Client name is required.")

        _assert_no_active_client_duplicate(db, client.pan, client.cin, client.gstin, exclude_id=client_id)
        client.updated_by = user["id"]
        client.updated_at = _now().isoformat()
        db.commit()
        db.refresh(client)
        result = _serialize_client(client)

    _log_history(user["id"], "CLIENT_UPDATE", {"client_id": result["id"], "display_name": result["display_name"]})
    return result

@api_router.delete("/clients/{client_id}")
async def delete_client(client_id: str, request: Request):
    user = await require_client_master_user(request)
    with _db() as db:
        client = db.query(Client).filter(Client.id == client_id, Client.is_deleted == 0).first()
        if not client:
            raise HTTPException(status_code=404, detail="Client not found.")
        client.is_deleted = 1
        client.updated_by = user["id"]
        client.updated_at = _now().isoformat()
        display_name = client.display_name
        db.commit()

    _log_history(user["id"], "CLIENT_DELETE", {"client_id": client_id, "display_name": display_name})
    return {"ok": True, "id": client_id}

@api_router.get("/history")
async def list_history(
    request: Request,
    email: Optional[str] = Query(default=None),
    start: Optional[str] = Query(default=None),
    end: Optional[str] = Query(default=None),
):
    user = await require_user(request)

    items = []
    with _db() as db:
        query = db.query(History, User.email).join(User, User.id == History.user_id)
        
        if user.get("role") != "admin":
            query = query.filter(History.user_id == user["id"])
        elif email:
            target = db.query(User).filter(User.email == email.strip().lower()).first()
            if target:
                query = query.filter(History.user_id == target.id)
            else:
                return {"items": []}

        if start:
            query = query.filter(History.timestamp >= start)
        if end:
            query = query.filter(History.timestamp <= end)

        for h_obj, u_email in query.order_by(History.timestamp.desc()).all():
            item = _to_dict(h_obj)
            item["email"] = u_email
            item["action_data"] = json.loads(item.get("action_data") or "{}")
            items.append(item)

    return {"items": items}

# -----------------------------
# Create universal certificate
# -----------------------------
@api_router.post("/certificates", response_model=UniversalCertificateStored)
async def create_certificate(payload: UniversalCertificateCreate, request: Request):
    # Enforce ownership: temp users only create under their own account
    user = await require_user(request)
    err = validate_common(payload)
    if err:
        raise HTTPException(status_code=400, detail=err)

    v = CATEGORY_VALIDATORS.get(payload.category)
    if v:
        err2 = v(payload)
        if err2:
            raise HTTPException(status_code=400, detail=err2)

    doc_data = {
        "client_id": payload.client_id,
        "category": payload.category,
        "certificate_type": payload.certificate_type,
        "entityType": payload.entityType,
        "identity": payload.identity.model_dump(),
        "meta": payload.meta.model_dump(),
        "ca": payload.ca.model_dump(),
        "data": payload.data.model_dump(),
    }
    doc = UniversalCertificateStored(**doc_data)

    stored = _serialize(doc.model_dump())
    cert_obj = Certificate(
        id=stored["id"],
        user_id=user["id"],
        created_by_info=user.get("email") or "system",
        category=stored["category"],
        certificate_type=stored["certificate_type"],
        entity_type=stored["entityType"],
        payload_json=json.dumps(stored, ensure_ascii=False, default=str),
        created_at=stored["created_at"],
        updated_at=stored["updated_at"]
    )
    with _db() as db:
        db.add(cert_obj)
        db.commit()
        
    _log_history(user["id"], "CERT_CREATE", {"cert_id": stored["id"], "category": stored["category"]})
    return doc

# -----------------------------
# List certificates (history)
# -----------------------------
@api_router.get("/certificates")
async def list_certificates(
    request: Request,
    page: Optional[int] = Query(default=None, ge=1),
    limit: Optional[int] = Query(default=None, ge=1, le=200),
):
    # Admin sees all, temporary users see only their own certificates
    user = await require_user(request)
    with _db() as db:
        query = db.query(Certificate.payload_json)
        if user.get("role") != "admin":
            query = query.filter(Certificate.user_id == user["id"])
        
        query = query.order_by(Certificate.created_at.desc())
        
        if page is not None and limit is not None:
            total = query.count()
            offset = (page - 1) * limit
            rows = query.offset(offset).limit(limit).all()
            certs = [json.loads(r[0]) for r in rows]
            return {"page": page, "limit": limit, "total": total, "items": certs}
        else:
            rows = query.all()
            certs = [json.loads(r[0]) for r in rows]
            return certs

# -----------------------------
# Get one certificate
# -----------------------------
@api_router.get("/certificates/{cert_id}")
async def get_certificate(cert_id: str, request: Request):
    user = await require_user(request)
    with _db() as db:
        row = db.query(Certificate.user_id, Certificate.payload_json).filter(Certificate.id == cert_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Certificate not found")
    if user.get("role") != "admin" and row.user_id != user["id"]:
        raise HTTPException(status_code=403, detail="Forbidden")
    return json.loads(row.payload_json)

# -----------------------------
# Update certificate (merge-safe)
# -----------------------------
@api_router.put("/certificates/{cert_id}")
async def update_certificate(cert_id: str, request: Request, payload: dict = Body(...)):
    user = await require_user(request)
    with _db() as db:
        cert = db.query(Certificate).filter(Certificate.id == cert_id).first()
        if not cert:
            raise HTTPException(status_code=404, detail="Certificate not found")
        
        is_admin = (user.get("role") or "").lower() == "admin"
        if not is_admin:
            if cert.user_id != user["id"]:
                raise HTTPException(status_code=403, detail="Forbidden")
            if not _can_user_manage_certificates(user):
                raise HTTPException(status_code=403, detail="Certificate manage permission denied.")

        existing = json.loads(cert.payload_json)
        updated = {**existing, **payload}
        updated["id"] = cert_id
        updated["created_at"] = existing.get("created_at", updated.get("created_at"))
        updated["certificate_type"] = updated.get("certificate_type") or existing.get("certificate_type")
        updated["updated_at"] = _now().isoformat()

        cert.payload_json = json.dumps(updated, ensure_ascii=False, default=str)
        cert.updated_at = updated["updated_at"]
        db.commit()

    _log_history(user["id"], "CERT_UPDATE", {"cert_id": cert_id})
    return updated

# -----------------------------
# Delete certificate
# -----------------------------
@api_router.delete("/certificates/{cert_id}")
async def delete_certificate(cert_id: str, request: Request):
    user = await require_user(request)
    with _db() as db:
        cert = db.query(Certificate).filter(Certificate.id == cert_id).first()
        if not cert:
            raise HTTPException(status_code=404, detail="Certificate not found")
        
        is_admin = (user.get("role") or "").lower() == "admin"
        if not is_admin:
            if cert.user_id != user["id"]:
                raise HTTPException(status_code=403, detail="Forbidden")
            if not _can_user_manage_certificates(user):
                raise HTTPException(status_code=403, detail="Certificate manage permission denied.")
        
        db.delete(cert)
        db.commit()
        
    _log_history(user["id"], "CERT_DELETE", {"cert_id": cert_id})
    return {"ok": True, "id": cert_id}

cors_uses_wildcard = "*" in CORS_ALLOW_ORIGINS
if IS_PRODUCTION and cors_uses_wildcard:
    raise RuntimeError("CORS_ORIGINS cannot contain '*' in production.")
trusted_hosts_use_wildcard = "*" in TRUSTED_ALLOWED_HOSTS
if IS_PRODUCTION and trusted_hosts_use_wildcard:
    raise RuntimeError("ALLOWED_HOSTS cannot contain '*' in production.")

app.add_middleware(
    TrustedHostMiddleware,
    allowed_hosts=TRUSTED_ALLOWED_HOSTS,
)

if FORCE_HTTPS:
    app.add_middleware(HTTPSRedirectMiddleware)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=not cors_uses_wildcard,
    allow_origins=CORS_ALLOW_ORIGINS,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type"],
)

# Wire router
# -----------------------------
app.include_router(api_router)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

